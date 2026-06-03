"""Text extraction from various document formats.

This module provides functionality for extracting text and structure
from PDFs, HTML, and plain text documents.
"""

import html
import re
import unicodedata

from thomas.marketplace.doc_processing._exceptions import ExtractionError
from thomas.marketplace.doc_processing._types import (
    BoundingBox,
    FontInfo,
    Footer,
    Header,
    Page,
    Paragraph,
    Sentence,
    Style,
    Table,
    TableCell,
    TextBlock,
    TextBlockType,
    Token,
)

# Complete, case-insensitive patterns for stripping <script>/<style> blocks.
# The tag body uses an unambiguous alternation (quoted attribute value | single
# non-">" char) so attribute values containing ">" are consumed without
# catastrophic backtracking, and the closing tag tolerates attributes/whitespace
# (e.g. "</script >"). re.IGNORECASE handles <SCRIPT>/<Style> variants.
_TAG_BODY = r'(?:"[^"]*"|\'[^\']*\'|[^>])*'
_SCRIPT_BLOCK_RE = re.compile(
    rf"<script{_TAG_BODY}>[\s\S]*?</script{_TAG_BODY}>",
    re.IGNORECASE,
)
_STYLE_BLOCK_RE = re.compile(
    rf"<style{_TAG_BODY}>[\s\S]*?</style{_TAG_BODY}>",
    re.IGNORECASE,
)


class TextExtractor:
    """Base class for text extraction from documents."""

    def extract(self, file_path: str) -> dict[str, any]:
        """Extract text and metadata from a document file.

        Supports multiple formats: PDF, DOCX, HTML, TXT, MD.
        Automatically detects format from file extension and content.

        Args:
            file_path: Path to document file (PDF, DOCX, HTML, TXT, MD)

        Returns:
            Dict with keys:
            - text: Extracted text content
            - metadata: Dict with document metadata
            - pages: Number of pages (for PDF, DOCX)
            - format: Detected format string

        Raises:
            ExtractionError: If file doesn't exist or extraction fails
        """
        from pathlib import Path

        try:
            path = Path(file_path)
            if not path.exists():
                raise ExtractionError(f"File not found: {file_path}")

            suffix = path.suffix.lower()
            text = ""
            pages = 0
            metadata: dict[str, any] = {"filename": path.name}

            # PDF extraction
            if suffix == ".pdf":
                try:
                    import pypdf

                    reader = pypdf.PdfReader(file_path)
                    pages = len(reader.pages)
                    text_parts = []
                    for page in reader.pages:
                        try:
                            text_parts.append(page.extract_text() or "")
                        except Exception:
                            pass
                    text = " ".join(text_parts)
                    metadata["format"] = "pdf"
                    metadata["page_count"] = pages

                except ImportError:
                    raise ExtractionError("pypdf not installed for PDF support")
                except Exception as e:
                    raise ExtractionError(f"PDF extraction failed: {e}")

            # DOCX extraction
            elif suffix == ".docx":
                try:
                    from docx import Document

                    doc = Document(file_path)
                    text_parts = [para.text for para in doc.paragraphs]
                    text = " ".join(text_parts)
                    pages = len(doc.sections) or 1
                    metadata["format"] = "docx"
                    metadata["paragraphs"] = len(doc.paragraphs)
                    metadata["tables"] = len(doc.tables)

                except ImportError:
                    raise ExtractionError("python-docx not installed for DOCX support")
                except Exception as e:
                    raise ExtractionError(f"DOCX extraction failed: {e}")

            # HTML extraction
            elif suffix in (".html", ".htm"):
                try:
                    text = path.read_text(encoding="utf-8")

                    # Remove script and style tags. Match the full tag (including
                    # attribute values that may contain ">") and the closing tag
                    # with optional attributes/whitespace, case-insensitively, so
                    # variants like <SCRIPT> or </script > cannot bypass the filter.
                    text = re.sub(_SCRIPT_BLOCK_RE, "", text)
                    text = re.sub(_STYLE_BLOCK_RE, "", text)

                    # Remove HTML tags
                    text = re.sub(r"<[^>]+>", " ", text)

                    # Decode HTML entities
                    text = html.unescape(text)

                    # Normalize whitespace
                    text = " ".join(text.split())

                    metadata["format"] = "html"
                    pages = 1

                except Exception as e:
                    raise ExtractionError(f"HTML extraction failed: {e}")

            # Plain text and Markdown
            elif suffix in (".txt", ".md"):
                try:
                    text = path.read_text(encoding="utf-8")
                    metadata["format"] = "markdown" if suffix == ".md" else "text"
                    pages = len(text.split("\n\n"))  # Paragraph count as proxy

                except Exception as e:
                    raise ExtractionError(f"Text file reading failed: {e}")

            else:
                # Try plain text as fallback
                try:
                    text = path.read_text(encoding="utf-8", errors="ignore")
                    metadata["format"] = "unknown"
                    pages = 1

                except Exception as e:
                    raise ExtractionError(f"File reading failed: {type(e).__name__}: {e}")

            # Normalize text
            text = PlainTextCleaner.clean(text)

            return {"text": text, "metadata": metadata, "pages": pages, "format": metadata.get("format", "unknown")}

        except ExtractionError:
            raise
        except Exception as e:
            raise ExtractionError(f"Extraction failed: {type(e).__name__}: {e}") from e


class PDFTextExtractor(TextExtractor):
    """Extract text from PDF documents (simulation).

    Simulates PDF text extraction including text block positioning,
    font information, and layout preservation.
    """

    def __init__(self) -> None:
        """Initialize PDF extractor."""
        self.pages: list[Page] = []

    def extract(self, source: str) -> str:
        """Extract text from PDF-like content.

        Args:
            source: PDF content string (simulated).

        Returns:
            Extracted text.

        Raises:
            ExtractionError: If extraction fails.
        """
        try:
            if not source:
                return ""

            text_parts: list[str] = []
            lines = source.split("\n")

            for line in lines:
                if line.strip():
                    text_parts.append(line)

            return " ".join(text_parts)
        except Exception as e:
            raise ExtractionError(f"PDF extraction failed: {e}")

    def extract_with_positions(self, source: str, page_width: float = 612.0, page_height: float = 792.0) -> list[Page]:
        """Extract text blocks with position information.

        Args:
            source: PDF content.
            page_width: Page width in points.
            page_height: Page height in points.

        Returns:
            List of Page objects with positioned text blocks.

        Raises:
            ExtractionError: If extraction fails.
        """
        try:
            pages: list[Page] = []
            lines = source.split("\n")

            text_blocks: list[TextBlock] = []
            current_y = 50.0
            page_number = 1

            for i, line in enumerate(lines):
                if not line.strip():
                    current_y += 12.0
                    continue

                # Simulate text block positioning
                font_size = 12.0
                is_title = current_y < 100 and i == 0
                is_heading = len(line) < 60 and current_y < 200

                if is_title:
                    block_type = TextBlockType.TITLE
                    font_size = 24.0
                elif is_heading:
                    block_type = TextBlockType.HEADING
                    font_size = 16.0
                else:
                    block_type = TextBlockType.PARAGRAPH
                    font_size = 12.0

                bbox = BoundingBox(
                    x0=50.0,
                    y0=current_y,
                    x1=page_width - 50.0,
                    y1=current_y + font_size + 4.0,
                    page_number=page_number,
                )

                font_info = FontInfo(
                    name="Helvetica",
                    size=font_size,
                    is_bold=is_title or is_heading,
                    is_italic=False,
                )

                style = Style(font=font_info)

                tokens = self._tokenize_line(line, bbox, style)

                sentences = self._create_sentences(tokens, line)

                paragraph = Paragraph(
                    sentences=sentences,
                    text=line,
                    block_type=block_type,
                    bbox=bbox,
                )

                text_block = TextBlock(
                    text=line,
                    bbox=bbox,
                    block_type=block_type,
                    paragraphs=[paragraph],
                    page_number=page_number,
                    confidence=0.95,
                )

                text_blocks.append(text_block)
                current_y += font_size + 6.0

                if current_y > page_height - 50:
                    page = Page(
                        page_number=page_number,
                        text="\n".join(b.text for b in text_blocks),
                        text_blocks=text_blocks,
                        width=page_width,
                        height=page_height,
                    )
                    pages.append(page)
                    text_blocks = []
                    current_y = 50.0
                    page_number += 1

            if text_blocks:
                page = Page(
                    page_number=page_number,
                    text="\n".join(b.text for b in text_blocks),
                    text_blocks=text_blocks,
                    width=page_width,
                    height=page_height,
                )
                pages.append(page)

            self.pages = pages
            return pages

        except Exception as e:
            raise ExtractionError(f"PDF extraction with positions failed: {e}")

    def _tokenize_line(self, line: str, bbox: BoundingBox, style: Style) -> list[Token]:
        """Tokenize a line into words.

        Args:
            line: Text line.
            bbox: Bounding box of line.
            style: Style information.

        Returns:
            List of Token objects.
        """
        tokens: list[Token] = []
        words = line.split()

        if not words:
            return tokens

        word_width = bbox.width() / len(words)
        current_x = bbox.x0

        for word in words:
            token_bbox = BoundingBox(
                x0=current_x,
                y0=bbox.y0,
                x1=current_x + word_width,
                y1=bbox.y1,
                page_number=bbox.page_number,
            )

            token = Token(
                text=word,
                bbox=token_bbox,
                style=style,
                page_number=bbox.page_number,
                confidence=0.95,
            )

            tokens.append(token)
            current_x += word_width

        return tokens

    def _create_sentences(self, tokens: list[Token], text: str) -> list[Sentence]:
        """Create sentences from tokens.

        Args:
            tokens: List of tokens.
            text: Original text.

        Returns:
            List of Sentence objects.
        """
        sentences: list[Sentence] = []

        # Simple sentence splitting on punctuation
        sentences_text = re.split(r"(?<=[.!?])\s+", text)

        token_idx = 0
        for sent_text in sentences_text:
            sent_tokens: list[Token] = []
            sent_word_count = len(sent_text.split())

            for _ in range(sent_word_count):
                if token_idx < len(tokens):
                    sent_tokens.append(tokens[token_idx])
                    token_idx += 1

            if sent_tokens:
                sentence = Sentence(
                    tokens=sent_tokens,
                    text=sent_text,
                    position=0,
                    page_number=tokens[0].page_number if tokens else 1,
                )
                sentences.append(sentence)

        return sentences


class HTMLTextExtractor(TextExtractor):
    """Extract text from HTML content.

    Removes tags, handles entities, and preserves document structure.
    """

    def __init__(self) -> None:
        """Initialize HTML extractor."""
        self.preserve_structure = True

    def extract(self, source: str) -> str:
        """Extract text from HTML.

        Args:
            source: HTML content.

        Returns:
            Extracted text.

        Raises:
            ExtractionError: If extraction fails.
        """
        try:
            # Remove script and style tags. Match the full tag (including
            # attribute values that may contain ">") and the closing tag with
            # optional attributes/whitespace, case-insensitively, so variants
            # like <SCRIPT> or </script > cannot bypass the filter.
            text = re.sub(_SCRIPT_BLOCK_RE, "", source)
            text = re.sub(_STYLE_BLOCK_RE, "", text)

            # Remove HTML tags
            text = re.sub(r"<[^>]+>", " ", text)

            # Decode HTML entities after stripping tags so escaped markup like
            # "&lt;tag&gt;" is preserved as text instead of being treated as HTML.
            text = html.unescape(text)

            # Normalize whitespace
            text = " ".join(text.split())

            return text

        except Exception as e:
            raise ExtractionError(f"HTML extraction failed: {e}")

    def extract_structured(self, source: str) -> dict[str, any]:
        """Extract structured content from HTML.

        Args:
            source: HTML content.

        Returns:
            Dictionary with extracted structure (headings, paragraphs, etc).

        Raises:
            ExtractionError: If extraction fails.
        """
        try:
            result: dict[str, any] = {
                "headings": [],
                "paragraphs": [],
                "lists": [],
                "tables": [],
                "links": [],
            }

            # Extract headings
            headings = re.findall(r"<h[1-6][^>]*>([^<]+)</h[1-6]>", source)
            result["headings"] = [html.unescape(h) for h in headings]

            # Extract paragraphs
            paragraphs = re.findall(r"<p[^>]*>([^<]+)</p>", source)
            result["paragraphs"] = [html.unescape(p) for p in paragraphs]

            # Extract list items
            list_items = re.findall(r"<li[^>]*>([^<]+)</li>", source)
            result["lists"] = [html.unescape(li) for li in list_items]

            # Extract links
            links = re.findall(r'<a[^>]*href="([^"]*)"[^>]*>([^<]+)</a>', source)
            result["links"] = [(url, html.unescape(text)) for url, text in links]

            return result

        except Exception as e:
            raise ExtractionError(f"HTML structured extraction failed: {e}")


class PlainTextCleaner:
    """Clean and normalize plain text content."""

    @staticmethod
    def normalize_unicode(text: str) -> str:
        """Normalize unicode characters.

        Args:
            text: Input text.

        Returns:
            Normalized text.
        """
        return unicodedata.normalize("NFKD", text)

    @staticmethod
    def remove_control_characters(text: str) -> str:
        """Remove control characters from text.

        Args:
            text: Input text.

        Returns:
            Text with control characters removed.
        """
        return "".join(char for char in text if unicodedata.category(char) != "Cc" or char in "\n\r\t")

    @staticmethod
    def normalize_whitespace(text: str) -> str:
        """Normalize whitespace.

        Args:
            text: Input text.

        Returns:
            Text with normalized whitespace.
        """
        # Replace multiple spaces with single space
        text = re.sub(r" +", " ", text)

        # Replace multiple newlines with double newline
        text = re.sub(r"\n\n+", "\n\n", text)

        # Remove trailing/leading whitespace
        return text.strip()

    @staticmethod
    def clean(text: str) -> str:
        """Full cleaning pipeline.

        Args:
            text: Input text.

        Returns:
            Cleaned text.
        """
        text = PlainTextCleaner.normalize_unicode(text)
        text = PlainTextCleaner.remove_control_characters(text)
        text = PlainTextCleaner.normalize_whitespace(text)
        return text


class TableExtractor:
    """Extract tables from text using column alignment analysis."""

    @staticmethod
    def detect_tables(text: str) -> list[Table]:
        """Detect tables in text using whitespace alignment.

        Detects aligned columns by analyzing consistent whitespace gaps.

        Args:
            text: Input text.

        Returns:
            List of detected tables.
        """
        lines = text.split("\n")
        tables: list[Table] = []

        if len(lines) < 2:
            return tables

        # Find lines with consistent spacing (potential table rows)
        table_start = 0
        in_table = False

        for i, line in enumerate(lines):
            # Heuristic: table rows have multiple space-separated fields
            fields = line.split()
            is_potential_row = len(fields) >= 2

            if is_potential_row and not in_table:
                in_table = True
                table_start = i
            elif not is_potential_row and in_table:
                # Extract table
                table_lines = lines[table_start:i]
                table = TableExtractor._parse_table_lines(table_lines)
                if table:
                    tables.append(table)
                in_table = False

        if in_table:
            table_lines = lines[table_start:]
            table = TableExtractor._parse_table_lines(table_lines)
            if table:
                tables.append(table)

        return tables

    @staticmethod
    def _parse_table_lines(lines: list[str]) -> Table | None:
        """Parse lines into table structure.

        Args:
            lines: Lines containing table data.

        Returns:
            Table object or None.
        """
        if not lines:
            return None

        # Detect column positions from first line
        first_line = lines[0]
        column_positions = TableExtractor._detect_columns(first_line)

        if len(column_positions) < 2:
            return None

        cells: list[TableCell] = []
        for row_idx, line in enumerate(lines):
            col_cells = TableExtractor._extract_cells_from_line(line, column_positions, row_idx)
            cells.extend(col_cells)

        if not cells:
            return None

        num_cols = len(column_positions)
        num_rows = len(lines)

        return Table(
            cells=cells,
            num_rows=num_rows,
            num_cols=num_cols,
            page_number=1,
        )

    @staticmethod
    def _detect_columns(line: str) -> list[int]:
        """Detect column positions in a line.

        Args:
            line: Text line.

        Returns:
            List of column start positions.
        """
        positions: list[int] = []
        in_space = True

        for i, char in enumerate(line):
            if char == " ":
                in_space = True
            else:
                if in_space:
                    positions.append(i)
                in_space = False

        return positions

    @staticmethod
    def _extract_cells_from_line(line: str, positions: list[int], row_idx: int) -> list[TableCell]:
        """Extract cells from a table line.

        Args:
            line: Table line.
            positions: Column positions.
            row_idx: Row index.

        Returns:
            List of TableCell objects.
        """
        cells: list[TableCell] = []

        for col_idx, pos in enumerate(positions):
            next_pos = positions[col_idx + 1] if col_idx + 1 < len(positions) else len(line)
            cell_text = line[pos:next_pos].strip()

            cell = TableCell(
                text=cell_text,
                row_index=row_idx,
                col_index=col_idx,
                is_header=(row_idx == 0),
            )
            cells.append(cell)

        return cells


class HeaderFooterDetector:
    """Detect and extract headers and footers from documents."""

    @staticmethod
    def detect_headers(pages: list[Page]) -> list[Header]:
        """Detect repeated headers across pages.

        Args:
            pages: List of pages.

        Returns:
            List of detected headers.
        """
        if len(pages) < 2:
            return []

        headers: list[Header] = []

        # Extract first line from each page
        first_lines: dict[str, int] = {}
        for page in pages:
            if page.text_blocks:
                first_block = page.text_blocks[0]
                first_text = first_block.text
                first_lines[first_text] = first_lines.get(first_text, 0) + 1

        # Repeated text is likely a header
        for text, count in first_lines.items():
            if count >= len(pages) * 0.5:  # Appears in at least 50% of pages
                for page in pages:
                    if page.text_blocks and page.text_blocks[0].text == text:
                        bbox = page.text_blocks[0].bbox
                        header = Header(
                            text=text,
                            bbox=bbox,
                            page_number=page.page_number,
                        )
                        headers.append(header)

        return headers

    @staticmethod
    def detect_footers(pages: list[Page]) -> list[Footer]:
        """Detect repeated footers across pages.

        Args:
            pages: List of pages.

        Returns:
            List of detected footers.
        """
        if len(pages) < 2:
            return []

        footers: list[Footer] = []

        # Extract last line from each page
        last_lines: dict[str, int] = {}
        for page in pages:
            if page.text_blocks:
                last_block = page.text_blocks[-1]
                last_text = last_block.text
                last_lines[last_text] = last_lines.get(last_text, 0) + 1

        # Repeated text is likely a footer
        for text, count in last_lines.items():
            if count >= len(pages) * 0.5:  # Appears in at least 50% of pages
                for page in pages:
                    if page.text_blocks and page.text_blocks[-1].text == text:
                        bbox = page.text_blocks[-1].bbox
                        footer = Footer(
                            text=text,
                            bbox=bbox,
                            page_number=page.page_number,
                        )
                        footers.append(footer)

        return footers


class FootnoteExtractor:
    """Extract footnotes and endnotes from documents."""

    @staticmethod
    def extract_footnotes(text: str) -> dict[str, list[str]]:
        """Extract footnotes from text.

        Detects common footnote markers and extracts note content.

        Args:
            text: Document text.

        Returns:
            Dictionary mapping footnote markers to content.
        """
        footnotes: dict[str, list[str]] = {
            "numbered": [],
            "symbols": [],
        }

        # Match numbered footnotes: [1], (1), etc.
        numbered = re.findall(r"\[(\d+)\]\s*([^\[\(]+?)(?=\[\d+\]|\([\d\*]+\)|$)", text)
        footnotes["numbered"] = [content.strip() for _, content in numbered]

        # Match symbol-based footnotes: [*], [†], etc.
        symbols = re.findall(r"\[([*†‡§¶])\]\s*([^\[\(]+?)(?=\[\d+\]|\[[\*†‡§¶]\]|$)", text)
        footnotes["symbols"] = [content.strip() for _, content in symbols]

        return footnotes
