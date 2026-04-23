"""Document ingestion — walk a folder, extract text, store in investigation DB.

Supports: PDF, text, HTML, JSON, CSV, images (stored for later LLM analysis).
No LLM calls here — pure Python text extraction.
"""

from __future__ import annotations

import csv
import hashlib
import json
import logging
import os
from collections.abc import Callable
from dataclasses import dataclass, field
from html.parser import HTMLParser
from pathlib import Path

from thomas.marketplace.investigation.store import Document, InvestigationStore

log = logging.getLogger(__name__)

# ---------------------------------------------------------------------------
# Supported file types
# ---------------------------------------------------------------------------

_TEXT_EXTS = {".txt", ".md", ".log", ".text", ".rtf"}
_PDF_EXTS = {".pdf"}
_HTML_EXTS = {".html", ".htm", ".mhtml"}
_JSON_EXTS = {".json"}
_CSV_EXTS = {".csv", ".tsv"}
_IMAGE_EXTS = {".png", ".jpg", ".jpeg", ".gif", ".bmp", ".webp", ".tiff"}
_DOCX_EXTS = {".docx"}
_EMAIL_EXTS = {".eml", ".msg"}

ALL_SUPPORTED = _TEXT_EXTS | _PDF_EXTS | _HTML_EXTS | _JSON_EXTS | _CSV_EXTS | _IMAGE_EXTS | _DOCX_EXTS | _EMAIL_EXTS


# ---------------------------------------------------------------------------
# Result
# ---------------------------------------------------------------------------


@dataclass
class IngestResult:
    total_files: int = 0
    ingested: int = 0
    skipped_existing: int = 0
    skipped_unsupported: int = 0
    failed: int = 0
    errors: list[tuple[str, str]] = field(default_factory=list)


# ---------------------------------------------------------------------------
# HTML stripper
# ---------------------------------------------------------------------------


class _HTMLTextExtractor(HTMLParser):
    def __init__(self):
        super().__init__()
        self._text: list[str] = []
        self._skip = False

    def handle_starttag(self, tag: str, attrs):
        if tag in ("script", "style", "noscript"):
            self._skip = True

    def handle_endtag(self, tag: str):
        if tag in ("script", "style", "noscript"):
            self._skip = False

    def handle_data(self, data: str):
        if not self._skip:
            stripped = data.strip()
            if stripped:
                self._text.append(stripped)

    def get_text(self) -> str:
        return "\n".join(self._text)


# ---------------------------------------------------------------------------
# Ingester
# ---------------------------------------------------------------------------


class DocumentIngester:
    """Walk a folder, extract text from files, store in InvestigationStore."""

    def __init__(self, store: InvestigationStore, *, max_text_chars: int = 100_000):
        self.store = store
        self.max_text_chars = max_text_chars

    def ingest_folder(
        self,
        folder_path: Path,
        case_id: str,
        *,
        recursive: bool = True,
        skip_existing: bool = True,
        on_progress: Callable[[str, int, int], None] | None = None,
    ) -> IngestResult:
        """Walk folder and ingest all supported files."""
        result = IngestResult()
        files = self._collect_files(folder_path, recursive=recursive)
        result.total_files = len(files)

        for i, fp in enumerate(files):
            if on_progress:
                on_progress(fp.name, i + 1, len(files))

            ext = fp.suffix.lower()
            if ext not in ALL_SUPPORTED:
                result.skipped_unsupported += 1
                continue

            try:
                file_hash = self._file_hash(fp)

                if skip_existing and self.store.document_exists(case_id, file_hash):
                    result.skipped_existing += 1
                    continue

                text, method, page_count = self._extract_text(fp)

                # Truncate extremely long text
                if text and len(text) > self.max_text_chars:
                    text = text[: self.max_text_chars] + "\n... (truncated)"

                doc = Document(
                    case_id=case_id,
                    file_path=str(fp.resolve()),
                    file_name=fp.name,
                    file_type=self._classify_type(ext),
                    file_size_bytes=fp.stat().st_size,
                    file_hash=file_hash,
                    extracted_text=text,
                    extraction_method=method,
                    page_count=page_count,
                )
                self.store.add_document(doc)
                result.ingested += 1

            except Exception as exc:
                log.warning("Failed to ingest %s: %s", fp, exc)
                result.failed += 1
                result.errors.append((str(fp), str(exc)))

        return result

    # -- File collection -----------------------------------------------------

    @staticmethod
    def _collect_files(folder: Path, *, recursive: bool = True) -> list[Path]:
        files: list[Path] = []
        if recursive:
            for root, _dirs, filenames in os.walk(folder):
                for fn in sorted(filenames):
                    fp = Path(root) / fn
                    if not fn.startswith("."):
                        files.append(fp)
        else:
            files = sorted(p for p in folder.iterdir() if p.is_file() and not p.name.startswith("."))
        return files

    # -- Text extraction -----------------------------------------------------

    def _extract_text(self, fp: Path) -> tuple[str | None, str, int | None]:
        """Returns (text, extraction_method, page_count)."""
        ext = fp.suffix.lower()

        if ext in _TEXT_EXTS:
            return self._extract_text_file(fp), "direct_read", None
        if ext in _PDF_EXTS:
            return self._extract_pdf(fp)
        if ext in _HTML_EXTS:
            return self._extract_html(fp), "html_strip", None
        if ext in _JSON_EXTS:
            return self._extract_json(fp), "json_parse", None
        if ext in _CSV_EXTS:
            return self._extract_csv(fp), "csv_read", None
        if ext in _IMAGE_EXTS:
            # Images: no text extraction here; analyzer will handle via LLM vision
            return None, "image_placeholder", None
        if ext in _DOCX_EXTS:
            return self._extract_docx(fp)
        if ext in _EMAIL_EXTS:
            return self._extract_email(fp), "email_parse", None

        return None, "unsupported", None

    @staticmethod
    def _extract_text_file(fp: Path) -> str:
        try:
            return fp.read_text(encoding="utf-8")
        except UnicodeDecodeError:
            return fp.read_text(encoding="latin-1")

    @staticmethod
    def _extract_pdf(fp: Path) -> tuple[str, str, int | None]:
        try:
            import pdfplumber
        except ImportError:
            # Try basic fallback
            return f"[PDF file: {fp.name} — install pdfplumber for text extraction]", "pdf_stub", None

        text_parts: list[str] = []
        page_count = 0
        with pdfplumber.open(fp) as pdf:
            page_count = len(pdf.pages)
            for page in pdf.pages[:200]:  # safety cap
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)

        return "\n\n".join(text_parts), "pdfplumber", page_count

    @staticmethod
    def _extract_html(fp: Path) -> str:
        try:
            raw = fp.read_text(encoding="utf-8")
        except UnicodeDecodeError:
            raw = fp.read_text(encoding="latin-1")

        extractor = _HTMLTextExtractor()
        extractor.feed(raw)
        return extractor.get_text()

    @staticmethod
    def _extract_json(fp: Path) -> str:
        try:
            raw = fp.read_text(encoding="utf-8")
            data = json.loads(raw)
            # Pretty-print for readability by LLM
            return json.dumps(data, ensure_ascii=False, indent=2)[:200_000]
        except (json.JSONDecodeError, UnicodeDecodeError):
            return fp.read_text(encoding="utf-8", errors="replace")

    @staticmethod
    def _extract_csv(fp: Path) -> str:
        lines: list[str] = []
        try:
            with open(fp, encoding="utf-8", newline="") as f:
                reader = csv.reader(f)
                for i, row in enumerate(reader):
                    if i >= 5000:  # safety cap
                        lines.append(f"... ({i}+ rows)")
                        break
                    lines.append(" | ".join(row))
        except Exception:
            return fp.read_text(encoding="utf-8", errors="replace")
        return "\n".join(lines)

    @staticmethod
    def _extract_docx(fp: Path) -> tuple[str, str, int | None]:
        try:
            import docx  # python-docx
        except ImportError:
            return (
                f"[Word file: {fp.name} — install python-docx for text extraction]",
                "docx_stub",
                None,
            )

        text_parts: list[str] = []
        doc = docx.Document(fp)

        # Extract paragraphs
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                text_parts.append(text)

        # Extract table contents
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    text_parts.append(" | ".join(cells))

        return "\n\n".join(text_parts), "python_docx", None

    @staticmethod
    def _extract_email(fp: Path) -> str:
        import email
        import email.policy

        try:
            raw = fp.read_bytes()
            msg = email.message_from_bytes(raw, policy=email.policy.default)
            parts: list[str] = []
            parts.append(f"From: {msg.get('From', '')}")
            parts.append(f"To: {msg.get('To', '')}")
            parts.append(f"Date: {msg.get('Date', '')}")
            parts.append(f"Subject: {msg.get('Subject', '')}")
            parts.append("")
            body = msg.get_body(preferencelist=("plain", "html"))
            if body:
                content = body.get_content()
                if isinstance(content, str):
                    parts.append(content)
            return "\n".join(parts)
        except Exception:
            return fp.read_text(encoding="utf-8", errors="replace")

    # -- Utilities -----------------------------------------------------------

    @staticmethod
    def _file_hash(fp: Path) -> str:
        h = hashlib.sha256()
        with open(fp, "rb") as f:
            while True:
                chunk = f.read(65536)
                if not chunk:
                    break
                h.update(chunk)
        return h.hexdigest()

    @staticmethod
    def _classify_type(ext: str) -> str:
        if ext in _PDF_EXTS:
            return "pdf"
        if ext in _TEXT_EXTS:
            return "text"
        if ext in _HTML_EXTS:
            return "html"
        if ext in _JSON_EXTS:
            return "json"
        if ext in _CSV_EXTS:
            return "csv"
        if ext in _IMAGE_EXTS:
            return "image"
        if ext in _EMAIL_EXTS:
            return "email"
        if ext in _DOCX_EXTS:
            return "docx"
        return "other"
